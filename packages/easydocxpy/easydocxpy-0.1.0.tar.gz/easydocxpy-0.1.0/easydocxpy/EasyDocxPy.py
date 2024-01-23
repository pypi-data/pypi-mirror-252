import logging
import os
from typing import Any, Optional

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text import paragraph


class EasyDocxPy:
    def __init__(self, titulo: str) -> None:
        # Configurar mensajes logs
        logging.basicConfig(
            level=logging.DEBUG,  # Establece el nivel de logging a DEBUG
            format="%(asctime)s - %(levelname)s - %(message)s",
            datefmt="%Y-%m-%d %H:%M:%S",  # Formato de fecha y hora
        )

        self.documento = Document()
        self.agregar_encabezado(titulo, 0)

    def agregar_encabezado(self, texto: str, nivel: int = 1) -> None:
        """
        Agrega un encabezado al documento de Word con el texto especificado y el nivel de encabezado deseado.

        :param texto: El texto que se va a agregar como encabezado.
        :type texto: str
        :param nivel: El nivel de encabezado (valor entre 0 y 9).
        :type nivel: int
        :raises ValueError: Si el nivel de encabezado está fuera del rango permitido.
        """
        # Comprobar que el nivel de encabezado se encuentra entre 0 y 9
        if not 0 <= nivel <= 9:
            mensaje_error = "El nivel de encabezado debe de estar entre 0 y 9."
            logging.error(mensaje_error)
            raise ValueError(mensaje_error)

        # Añadir encabezado
        self.documento.add_heading(texto, level=nivel)

    def agregar_parrafo(
        self, texto: str = "", centrado: bool = False, espacio_despues: int = 12
    ) -> Any:
        """
        Agrega un párrafo al documento de Word con el texto especificado, la alineación deseada y un espacio después del párrafo.

        :param texto: El texto que se va a agregar al párrafo.
        :type texto: str
        :param centrado: Indica si se debe centrar el párrafo (predeterminado es False).
        :type centrado: bool
        :param espacio_despues: La cantidad de espacio en puntos (Pt) que se agregará después del párrafo (predeterminado es 12 Pt).
        :type espacio_despues: int
        :return: El párrafo creado.
        :rtype: docx.text.paragraph.Paragraph
        """
        # Añadir texto al nuevo párrafo creado
        parrafo: paragraph = self.documento.add_paragraph(texto)

        # Configurar el espacio después del párrafo
        formato_parrafo = parrafo.paragraph_format
        formato_parrafo.space_after = Pt(espacio_despues)

        if centrado:
            parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        return parrafo

    def agregar_texto_al_parrafo(
        self, parrafo, texto, negrita=False, cursiva=False, subrayado=False
    ) -> Any:
        """
        Agrega texto a un párrafo existente y aplica formato (negrita, cursiva, subrayado) según sea necesario.

        :param parrafo: El párrafo al que se agregará el texto.
        :type parrafo: docx.text.paragraph.Paragraph
        :param texto: El texto que se va a agregar al párrafo.
        :type texto: str
        :param negrita: Indica si el texto debe aparecer en negrita (predeterminado es False).
        :type negrita: bool
        :param cursiva: Indica si el texto debe aparecer en cursiva (predeterminado es False).
        :type cursiva: bool
        :param subrayado: Indica si el texto debe aparecer subrayado (predeterminado es False).
        :type subrayado: bool
        :return: El fragmento de texto creado.
        :rtype: docx.text.run.Run
        """
        # Añadir texto
        run = parrafo.add_run(texto)
        # Aplicar formato según parámetros de entrada
        run.bold = negrita
        run.italic = cursiva
        run.underline = subrayado
        return run

    def agregar_imagen(
        self,
        ruta_imagen: str,
        parrafo=None,
        centrado: bool = True,
        ancho: int = 6,
    ) -> None:
        """
        Agrega una imagen al documento de Word.

        :param ruta_imagen: La ruta de la imagen que se va a agregar al documento.
        :type ruta_imagen: str
        :param parrafo: El párrafo al que se agregará la imagen (opcional).
        :type parrafo: docx.text.paragraph.Paragraph
        :param centrado: Indica si se debe centrar la imagen en el párrafo (predeterminado es True).
        :type centrado: bool
        :param ancho: El ancho de la imagen en pulgadas (predeterminado es 6).
        :type ancho: int
        :param heading: Texto de encabezado opcional para la imagen (predeterminado es None).
        :type heading: str
        :raises Exception: Si ocurre un error al insertar la imagen.
        """
        # Si no se proporciona un párrafo, se agrega uno nuevo
        if parrafo is None:
            parrafo = self.agregar_parrafo(centrado=centrado)

        # Crear un fragmento de texto en el párrafo actual
        run = parrafo.add_run()

        try:
            # Agregar la imagen al párrafo
            run.add_picture(ruta_imagen, width=Inches(ancho))
        except Exception as e:
            mensaje_error: str = f"Error al insertar la imagen: {e}"
            logging.error(mensaje_error)
            # Si la ruta de la imagen no existe, se asigna una imagen predeterminada
            run.add_picture("assets/no-fotos.png", width=Inches(1.5))
            # Agregar mensaje del error
            self.agregar_parrafo(texto=mensaje_error, centrado=True)
            # raise Exception(mensaje_error) from None
        else:
            # Si se especifica centrado, alinear el párrafo al centro
            if centrado:
                parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def agregar_dataframe(self, dataframe: pd.DataFrame) -> None:
        """
        Agrega un DataFrame a un documento de Word como una tabla, formateando los títulos de las columnas
        y los datos de las celdas.

        :param dataframe: El DataFrame de pandas que se va a agregar como una tabla.
        :type dataframe: pd.DataFrame
        """
        # Agregar una tabla con el número correcto de filas y columnas
        rows: int = len(dataframe) + 1
        cols: int = len(dataframe.columns)
        tabla: Table = self.documento.add_table(rows=rows, cols=cols)

        # Establecer la alineación vertical y horizontal, y aplicar negrita a los títulos de las columnas
        for j, col in enumerate(dataframe.columns):
            cell = tabla.cell(0, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(str(col))
            run.bold = True

        # Establecer la alineación vertical y horizontal para los datos del DataFrame
        for i, fila in enumerate(dataframe.values):
            for j, valor in enumerate(fila):
                cell = tabla.cell(i + 1, j)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run(str(valor))

    def guardar_documento(self, nombre_archivo: str) -> None:
        """
        Guarda el documento de Word con un nombre de archivo especificado.

        :param nombre_archivo: El nombre de archivo para guardar el documento (sin extensión).
        :type nombre_archivo: str
        """
        self.documento.save(nombre_archivo + ".docx")


if __name__ == "__main__":
    # Uso de la clase
    reporte = EasyDocxPy("Mi Reporte")
    reporte.agregar_parrafo("Aquí va el texto del reporte.")
    reporte.agregar_imagen("captura.png")

    data_dict = {
        "ID": [1, 2, 3],
        "Nombre": ["Alice", "Bob", "Max"],
        "Edad": [25, 30, 35],
    }
    data = pd.DataFrame(data_dict)

    reporte.agregar_dataframe(data)
    reporte.guardar_documento("ReporteEjemplo")
